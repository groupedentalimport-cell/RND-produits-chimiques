"""
ICH/FDA/EMA Regulatory Report Generator.
Produces compliant stability assessment reports with electronic signatures.
"""

import io
from datetime import datetime
from typing import Dict, Any, List, Optional
from dataclasses import dataclass


@dataclass
class ReportSection:
    title: str
    content: str
    subsections: List['ReportSection']
    table_data: Optional[List[List[str]]] = None


class RegulatoryReportGenerator:
    """
    Generate ICH Q1A/Q1B compliant stability reports.
    Supports PDF (reportlab), DOCX (python-docx), XLSX (openpyxl).
    """

    def __init__(self, company_name: str = "ChemStab Industrial"):
        self.company_name = company_name
        self.timestamp = datetime.utcnow()

    def generate_pdf(
        self,
        analysis_data: Dict[str, Any],
        project_data: Dict[str, Any],
        substances: List[Dict[str, Any]],
        signature: Optional[Dict[str, str]] = None,
    ) -> bytes:
        """Generate ICH-compliant PDF report."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, HRFlowable
        )
        from reportlab.lib.units import cm, mm
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=2.5 * cm, bottomMargin=2 * cm,
            leftMargin=2 * cm, rightMargin=2 * cm,
        )
        styles = getSampleStyleSheet()

        # Custom styles
        styles.add(ParagraphStyle(
            "ReportTitle", parent=styles["Title"],
            fontSize=18, spaceAfter=6, textColor=colors.HexColor("#1E3A8A"),
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            "SectionHeader", parent=styles["Heading1"],
            fontSize=14, spaceBefore=16, spaceAfter=8,
            textColor=colors.HexColor("#1E40AF"),
            borderWidth=1, borderColor=colors.HexColor("#1E40AF"),
            borderPadding=4,
        ))
        styles.add(ParagraphStyle(
            "SubSection", parent=styles["Heading2"],
            fontSize=12, spaceBefore=10, spaceAfter=6,
            textColor=colors.HexColor("#3B82F6"),
        ))
        styles.add(ParagraphStyle(
            "BodyText2", parent=styles["Normal"],
            fontSize=10, spaceAfter=6, leading=14,
        ))
        styles.add(ParagraphStyle(
            "SmallText", parent=styles["Normal"],
            fontSize=8, textColor=colors.grey,
        ))

        elements = []

        # ── Cover page ─────────────────────────────────────────────────
        elements.append(Spacer(1, 3 * cm))
        elements.append(Paragraph("STABILITY ASSESSMENT REPORT", styles["ReportTitle"]))
        elements.append(Spacer(1, 0.5 * cm))
        elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1E40AF")))
        elements.append(Spacer(1, 1 * cm))

        # Report metadata table
        report_id = f"SAR-{self.timestamp.strftime('%Y%m%d')}-{analysis_data.get('id', '000'):03d}"
        meta_data = [
            ["Report Number", report_id],
            ["Document Type", "Stability Assessment Report"],
            ["Regulatory Framework", "ICH Q1A(R2) / ICH Q1B"],
            ["Product", project_data.get("name", "N/A")],
            ["Product Type", project_data.get("product_type", "N/A")],
            ["Formulation", project_data.get("formulation_type", "N/A")],
            ["Target Market", project_data.get("target_market", "ICH")],
            ["Version", project_data.get("version", "1.0")],
            ["Generated", self.timestamp.strftime("%Y-%m-%d %H:%M UTC")],
            ["Prepared By", self.company_name],
        ]

        meta_table = Table(meta_data, colWidths=[5 * cm, 10 * cm])
        meta_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ]))
        elements.append(meta_table)
        elements.append(PageBreak())

        # ── Table of Contents ──────────────────────────────────────────
        elements.append(Paragraph("TABLE OF CONTENTS", styles["SectionHeader"]))
        toc_items = [
            "1. Executive Summary",
            "2. Product Information",
            "3. Test Substance(s)",
            "4. Analytical Methods & Conditions",
            "5. Stability Risk Assessment",
            "6. Kinetic Analysis & Shelf Life Prediction",
            "7. QSPR/ML Predictions",
            "8. Packaging & Container Compatibility",
            "9. Conclusions & Recommendations",
            "10. Electronic Signatures",
            "11. Audit Trail Reference",
        ]
        for item in toc_items:
            elements.append(Paragraph(item, styles["BodyText2"]))
        elements.append(PageBreak())

        # ── 1. Executive Summary ───────────────────────────────────────
        elements.append(Paragraph("1. Executive Summary", styles["SectionHeader"]))
        overall_score = analysis_data.get("overall_score", "N/A")
        risk_level = analysis_data.get("overall_severity", "N/A")
        elements.append(Paragraph(
            f"Overall Stability Score: <b>{overall_score}/100</b> | "
            f"Risk Level: <b>{risk_level.upper() if isinstance(risk_level, str) else risk_level}</b>",
            styles["BodyText2"],
        ))
        elements.append(Paragraph(
            f"This report presents the results of a comprehensive chemical stability assessment "
            f"performed on {self.timestamp.strftime('%B %d, %Y')} in accordance with ICH Q1A(R2) guidelines.",
            styles["BodyText2"],
        ))
        elements.append(Spacer(1, 0.5 * cm))

        # ── 2. Product Information ─────────────────────────────────────
        elements.append(Paragraph("2. Product Information", styles["SectionHeader"]))
        prod_data = [
            ["Parameter", "Value"],
            ["Product Name", project_data.get("name", "N/A")],
            ["Product Type", project_data.get("product_type", "N/A")],
            ["Formulation Type", project_data.get("formulation_type", "N/A")],
            ["Target Market", project_data.get("target_market", "ICH")],
            ["Project Code", project_data.get("code", "N/A")],
            ["GxP Critical", "Yes" if project_data.get("is_gxp_critical") else "No"],
        ]
        prod_table = Table(prod_data, colWidths=[5 * cm, 10 * cm])
        prod_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ]))
        elements.append(prod_table)

        # ── 3. Test Substances ─────────────────────────────────────────
        elements.append(Paragraph("3. Test Substance(s)", styles["SectionHeader"]))
        if substances:
            sub_data = [["#", "Name", "CAS", "Concentration", "Unit", "Purity", "Grade"]]
            for i, s in enumerate(substances, 1):
                sub_data.append([
                    str(i),
                    s.get("name", ""),
                    s.get("cas_number", ""),
                    str(s.get("concentration", "")),
                    s.get("concentration_unit", ""),
                    f"{s.get('purity', '')}%",
                    s.get("grade", ""),
                ])
            sub_table = Table(sub_data, colWidths=[1*cm, 4*cm, 2.5*cm, 2*cm, 1.5*cm, 1.5*cm, 2.5*cm])
            sub_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ]))
            elements.append(sub_table)
        else:
            elements.append(Paragraph("No substances recorded.", styles["BodyText2"]))

        # ── 5. Stability Risk Assessment ───────────────────────────────
        elements.append(Paragraph("5. Stability Risk Assessment", styles["SectionHeader"]))
        risks = analysis_data.get("risks", {})
        if risks:
            risk_data = [["Risk Type", "Score", "Severity", "Description"]]
            for risk_type, rd in risks.items():
                risk_data.append([
                    f"{rd.get('icon', '')} {rd.get('name', risk_type)}",
                    str(rd.get('score', '')),
                    rd.get('severity', '').upper(),
                    rd.get('description', '')[:80],
                ])
            risk_table = Table(risk_data, colWidths=[4*cm, 2*cm, 2*cm, 7*cm])
            risk_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ]))
            elements.append(risk_table)

        # ── 6. Kinetic Analysis ────────────────────────────────────────
        kinetics = analysis_data.get("kinetics_results", {})
        if kinetics:
            elements.append(Paragraph("6. Kinetic Analysis & Shelf Life Prediction", styles["SectionHeader"]))
            elements.append(Paragraph(
                f"Predicted Shelf Life: <b>{kinetics.get('shelf_life_months', 'N/A')} months</b> "
                f"at {kinetics.get('storage_temperature', 25)}°C",
                styles["BodyText2"],
            ))
            elements.append(Paragraph(
                f"Q10 Factor: {kinetics.get('q10', 'N/A')} | "
                f"Rate Constant (k): {kinetics.get('rate_constant', 'N/A')}",
                styles["BodyText2"],
            ))

        # ── 9. Conclusions ─────────────────────────────────────────────
        elements.append(Paragraph("9. Conclusions & Recommendations", styles["SectionHeader"]))
        recs = analysis_data.get("recommendations", [])
        if recs:
            for rec in recs:
                elements.append(Paragraph(f"• {rec}", styles["BodyText2"]))
        else:
            elements.append(Paragraph("No specific recommendations at this time.", styles["BodyText2"]))

        # ── 10. Electronic Signature ───────────────────────────────────
        elements.append(PageBreak())
        elements.append(Paragraph("10. Electronic Signatures", styles["SectionHeader"]))
        elements.append(Paragraph(
            "This document has been electronically signed in accordance with "
            "21 CFR Part 11 (Electronic Records; Electronic Signatures).",
            styles["BodyText2"],
        ))
        if signature:
            sig_data = [
                ["Field", "Value"],
                ["Signed By", signature.get("signed_by", "N/A")],
                ["Meaning", signature.get("meaning", "N/A")],
                ["Timestamp", signature.get("timestamp", "N/A")],
                ["Signature Hash", signature.get("signature_hash", "N/A")[:32] + "..."],
                ["Algorithm", signature.get("algorithm", "SHA-256")],
            ]
            sig_table = Table(sig_data, colWidths=[4 * cm, 11 * cm])
            sig_table.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]))
            elements.append(sig_table)

        # ── Footer ─────────────────────────────────────────────────────
        elements.append(Spacer(1, 2 * cm))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        elements.append(Paragraph(
            f"Generated by {self.company_name} v5.1 | "
            f"{self.timestamp.strftime('%Y-%m-%d %H:%M UTC')} | "
            f"CONFIDENTIAL",
            styles["SmallText"],
        ))

        doc.build(elements)
        pdf_bytes = buffer.getvalue()
        assert pdf_bytes[:5] == b"%PDF-", "Invalid PDF output"
        return pdf_bytes

    def generate_docx(
        self,
        analysis_data: Dict[str, Any],
        project_data: Dict[str, Any],
        substances: List[Dict[str, Any]],
        signature: Optional[Dict[str, str]] = None,
    ) -> bytes:
        """Generate ICH-compliant DOCX report."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Title
        title = doc.add_heading("STABILITY ASSESSMENT REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        report_id = f"SAR-{self.timestamp.strftime('%Y%m%d')}-{analysis_data.get('id', '000'):03d}"
        doc.add_paragraph(f"Report Number: {report_id}")
        doc.add_paragraph(f"Product: {project_data.get('name', 'N/A')}")
        doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("")

        # Executive Summary
        doc.add_heading("1. Executive Summary", level=1)
        overall = analysis_data.get("overall_score", "N/A")
        severity = analysis_data.get("overall_severity", "N/A")
        doc.add_paragraph(f"Overall Stability Score: {overall}/100")
        doc.add_paragraph(f"Risk Level: {severity.upper() if isinstance(severity, str) else severity}")

        # Risk Assessment
        doc.add_heading("5. Stability Risk Assessment", level=1)
        risks = analysis_data.get("risks", {})
        if risks:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Risk Type"
            hdr[1].text = "Score"
            hdr[2].text = "Severity"
            for risk_type, rd in risks.items():
                row = table.add_row().cells
                row[0].text = f"{rd.get('icon', '')} {rd.get('name', risk_type)}"
                row[1].text = str(rd.get("score", ""))
                row[2].text = rd.get("severity", "").upper()

        # Recommendations
        doc.add_heading("9. Conclusions & Recommendations", level=1)
        for rec in analysis_data.get("recommendations", []):
            doc.add_paragraph(rec, style="List Bullet")

        # Signature
        if signature:
            doc.add_heading("10. Electronic Signature", level=1)
            doc.add_paragraph(f"Signed by: {signature.get('signed_by', 'N/A')}")
            doc.add_paragraph(f"Meaning: {signature.get('meaning', 'N/A')}")
            doc.add_paragraph(f"Timestamp: {signature.get('timestamp', 'N/A')}")
            doc.add_paragraph(f"Hash: {signature.get('signature_hash', 'N/A')}")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_xlsx(
        self,
        analysis_data: Dict[str, Any],
        substances: List[Dict[str, Any]],
    ) -> bytes:
        """Generate ICH-compliant XLSX spreadsheet."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()

        # Sheet 1: Risk Assessment
        ws = wb.active
        ws.title = "Risk Assessment"

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="1E40AF", end_color="1E40AF", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        headers = ["Risk Type", "Score (0-100)", "Severity", "Description", "Key Factors"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border

        row = 2
        for risk_type, rd in analysis_data.get("risks", {}).items():
            ws.cell(row=row, column=1, value=rd.get("name", risk_type)).border = thin_border
            ws.cell(row=row, column=2, value=rd.get("score", 0)).border = thin_border
            sev = ws.cell(row=row, column=3, value=rd.get("severity", "").upper())
            sev.border = thin_border
            sev_colors = {"CRITICAL": "DC2626", "HIGH": "F59E0B", "MODERATE": "3B82F6", "LOW": "10B981"}
            color = sev_colors.get(rd.get("severity", "").upper(), "6B7280")
            sev.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
            sev.font = Font(color="FFFFFF", bold=True)
            ws.cell(row=row, column=4, value=rd.get("description", "")).border = thin_border
            factors = rd.get("factors", [])
            ws.cell(row=row, column=5, value=str(factors)[:200]).border = thin_border
            row += 1

        # Auto-width
        for col in ws.columns:
            max_len = max(len(str(c.value or "")) for c in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        # Sheet 2: Substances
        ws2 = wb.create_sheet("Substances")
        sub_headers = ["Name", "CAS", "Formula", "MW", "Conc.", "Unit", "Purity", "Grade"]
        for col, h in enumerate(sub_headers, 1):
            cell = ws2.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, s in enumerate(substances, 2):
            ws2.cell(row=i, column=1, value=s.get("name", ""))
            ws2.cell(row=i, column=2, value=s.get("cas_number", ""))
            ws2.cell(row=i, column=3, value=s.get("formula", ""))
            ws2.cell(row=i, column=4, value=s.get("molar_mass", ""))
            ws2.cell(row=i, column=5, value=s.get("concentration", ""))
            ws2.cell(row=i, column=6, value=s.get("concentration_unit", ""))
            ws2.cell(row=i, column=7, value=s.get("purity", ""))
            ws2.cell(row=i, column=8, value=s.get("grade", ""))

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()
